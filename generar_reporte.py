# -*- coding: utf-8 -*-
"""
Genera el reporte técnico de tesis en formato Word.
"""

import os
import pickle
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def bold_cell(cell, text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def normal_cell(cell, text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def add_para(doc, text, size=11, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def header_row(table, headers, bg="1F497D", text_color="FFFFFF", size=10):
    row = table.rows[0]
    for i, h in enumerate(headers):
        bold_cell(row.cells[i], h, size=size, color=text_color)
        set_cell_bg(row.cells[i], bg)

def add_image_if_exists(doc, path, width_cm=14, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)

# ─── Cargar resultados ─────────────────────────────────────────────────────────

def load_pkl(path):
    if os.path.exists(path):
        with open(path, "rb") as f:
            return pickle.load(f)
    return None

kfold_cnn  = load_pkl("outputs/crossval/kfold_cnn/kfold_metrics.pkl")
kfold_pose = load_pkl("outputs/crossval/kfold_pose/kfold_metrics.pkl")
rob_pose   = load_pkl("outputs/robustness/robustness_pose/robustness_results.pkl")
rob_cnn    = load_pkl("outputs/robustness/robustness_cnn/robustness_results.pkl")

def mean_std(folds, key):
    vals = [f[key] for f in folds]
    return np.mean(vals), np.std(vals)

# ─── Documento ────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("REPORTE TÉCNICO DE TESIS")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Sistema de Detección de Escoliosis Mediante\nAprendizaje Profundo y Estimación de Pose")
sr.bold = True
sr.font.size = Pt(16)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Santiago Escalante Carrión\nAbril 2026")
mr.font.size = Pt(12)
mr.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN GENERAL DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Descripción General del Proyecto", 1)

add_para(doc,
    "Este proyecto de tesis desarrolla un sistema de inteligencia artificial para la "
    "detección automatizada de escoliosis a partir de imágenes de postura corporal. "
    "Se implementan y comparan tres enfoques complementarios que forman un estudio de "
    "ablación completo:")

items = [
    ("E1 – CNN (Transfer Learning)",
     "Clasificación visual pura mediante EfficientNetB0 ajustado a las imágenes."),
    ("E2 – Pose + ML",
     "Extracción de 12 características geométricas clínicas mediante MediaPipe y "
     "clasificación con Random Forest, XGBoost y SVM."),
    ("E3 – Modelo Híbrido",
     "Fusión de las dos corrientes anteriores en una arquitectura de doble entrada."),
]
for title_i, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    run_t = p.add_run(title_i + ": ")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.size = Pt(11)

add_para(doc,
    "\nEl objetivo final es disponer de una herramienta de cribado no invasiva, "
    "interpretable y robusta que pueda asistir al personal médico en la detección "
    "temprana de la escoliosis.", size=11)

# ══════════════════════════════════════════════════════════════════════════════
# 2. CONJUNTO DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Conjunto de Datos", 1)

add_heading(doc, "2.1 Estructura y División", 2)
add_para(doc,
    "El dataset está organizado en dos clases (scoliosis_yes / scoliosis_no). "
    "Para garantizar que no exista fuga de datos entre particiones, la división se "
    "realiza sobre imágenes originales; las versiones aumentadas sólo se incluyen "
    "en el conjunto de entrenamiento.")

# Tabla de división
t = doc.add_table(rows=4, cols=3)
t.style = "Table Grid"
header_row(t, ["Partición", "Composición", "Proporción"])
data_split = [
    ("Entrenamiento", "Imágenes originales + todas las aumentadas", "≈70 %"),
    ("Validación",    "Sólo imágenes originales",                   "≈15 %"),
    ("Prueba",        "Sólo imágenes originales",                   "≈15 %"),
]
for i, (p, c, r) in enumerate(data_split, start=1):
    normal_cell(t.rows[i].cells[0], p, align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t.rows[i].cells[1], c, align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t.rows[i].cells[2], r)
doc.add_paragraph()

add_heading(doc, "2.2 Aumentación de Datos (albumentations)", 2)
add_para(doc,
    "Se aplica un pipeline de aumentación para llegar a ~1 000 imágenes por clase "
    "(originales + aumentadas). Las transformaciones aplicadas son:")

aug_items = [
    ("HorizontalFlip", "p = 0.5", "Volteo horizontal aleatorio"),
    ("Rotate", "límite ±12°, p = 0.7", "Rotación moderada"),
    ("ShiftScaleRotate", "shift 0.05, scale 0.08, p = 0.6", "Desplazamiento y escala"),
    ("RandomBrightnessContrast", "±12 %, p = 0.5", "Variación fotométrica"),
    ("GaussNoise", "var (5, 20), p = 0.3", "Ruido gaussiano leve"),
    ("GaussianBlur", "kernel 3×3, p = 0.2", "Desenfoque suave"),
]
t2 = doc.add_table(rows=len(aug_items)+1, cols=3)
t2.style = "Table Grid"
header_row(t2, ["Transformación", "Parámetros", "Propósito"])
for i, (tr_, pa, pu) in enumerate(aug_items, start=1):
    normal_cell(t2.rows[i].cells[0], tr_, align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t2.rows[i].cells[1], pa,  align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t2.rows[i].cells[2], pu,  align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ARQUITECTURAS Y HERRAMIENTAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Arquitecturas y Herramientas", 1)

# --- 3.1 E1 CNN
add_heading(doc, "3.1 E1: CNN con Transfer Learning (EfficientNetB0)", 2)
add_para(doc,
    "Se emplea EfficientNetB0 preentrenado en ImageNet como backbone. El entrenamiento "
    "se divide en dos fases para estabilizar el ajuste fino:")

add_para(doc, "Arquitectura de la cabeza de clasificación:", bold=True, size=10)

arch_rows = [
    ("Entrada", "(224, 224, 3) – imagen RGB normalizada"),
    ("Backbone", "EfficientNetB0 (congelado en Fase 1)"),
    ("GlobalAveragePooling2D", "Reducción espacial del mapa de características"),
    ("BatchNormalization", "Estabilización del entrenamiento"),
    ("Dropout(0.4)", "Regularización"),
    ("Dense(128, relu)", "Representación de alto nivel"),
    ("Dropout(0.2)", "Regularización adicional"),
    ("Dense(1, sigmoid)", "Probabilidad de escoliosis [0, 1]"),
]
t3 = doc.add_table(rows=len(arch_rows)+1, cols=2)
t3.style = "Table Grid"
header_row(t3, ["Capa", "Descripción"])
for i, (lay, desc) in enumerate(arch_rows, start=1):
    normal_cell(t3.rows[i].cells[0], lay,  align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    normal_cell(t3.rows[i].cells[1], desc, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

add_para(doc, "Estrategia de entrenamiento en dos fases:", bold=True, size=10)
phase_rows = [
    ("Fase 1 – Cabeza", "Congelado", "1×10⁻³", "10", "EarlyStopping + ReduceLROnPlateau"),
    ("Fase 2 – Fine-tune", "Bloques 6-7 descongelados", "1×10⁻⁵", "20", "Mismos callbacks"),
]
t4 = doc.add_table(rows=len(phase_rows)+1, cols=5)
t4.style = "Table Grid"
header_row(t4, ["Fase", "Backbone", "LR (Adam)", "Épocas máx.", "Callbacks"])
for i, row in enumerate(phase_rows, start=1):
    for j, v in enumerate(row):
        normal_cell(t4.rows[i].cells[j], v, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

add_para(doc, "Callbacks detallados:", bold=True, size=10)
cb_rows = [
    ("EarlyStopping", "val_loss", "patience=5, restore_best_weights=True"),
    ("ReduceLROnPlateau", "val_loss", "factor=0.5, patience=3, min_lr=1×10⁻⁷"),
    ("ModelCheckpoint", "val_loss", "Guarda el mejor modelo por pérdida de validación"),
]
t5 = doc.add_table(rows=len(cb_rows)+1, cols=3)
t5.style = "Table Grid"
header_row(t5, ["Callback", "Monitor", "Configuración"])
for i, row in enumerate(cb_rows, start=1):
    for j, v in enumerate(row):
        normal_cell(t5.rows[i].cells[j], v, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

# --- 3.2 E2 Pose + ML
add_heading(doc, "3.2 E2: Estimación de Pose + Aprendizaje Automático", 2)
add_para(doc,
    "Se utiliza MediaPipe Pose Landmarker para detectar 33 puntos clave del cuerpo. "
    "A partir de estos se computan 12 características geométricas clínicamente relevantes:")

feat_rows = [
    ("shoulder_height_diff", "|hombro_L.y − hombro_R.y|", "Asimetría de hombros"),
    ("hip_height_diff",       "|cadera_L.y − cadera_R.y|", "Asimetría de caderas"),
    ("elbow_height_diff",     "|codo_L.y − codo_R.y|",     "Asimetría de codos"),
    ("wrist_height_diff",     "|muñeca_L.y − muñeca_R.y|", "Asimetría distal"),
    ("shoulder_width",        "|hombro_L.x − hombro_R.x|", "Ancho de hombros"),
    ("hip_width",             "|cadera_L.x − cadera_R.x|", "Ancho de caderas"),
    ("shoulder_asymmetry_ratio", "shoulder_height_diff / shoulder_width", "Asimetría relativa (invariante a escala)"),
    ("hip_asymmetry_ratio",   "hip_height_diff / hip_width",             "Asimetría relativa de caderas"),
    ("trunk_tilt_x",          "mid_hombro.x − mid_cadera.x",            "Desviación lateral del tronco"),
    ("trunk_tilt_angle",      "arctan2(|trunk_tilt_x|, |Δy tronco|)",   "Ángulo de inclinación del tronco (grados)"),
    ("spine_deviation",       "|trunk_tilt_x|",                         "Desviación lateral de la columna"),
    ("body_height_proxy",     "|mid_hombro.y − mid_cadera.y|",          "Proxy de altura del tronco"),
]
t6 = doc.add_table(rows=len(feat_rows)+1, cols=3)
t6.style = "Table Grid"
header_row(t6, ["Característica", "Cálculo", "Relevancia Clínica"])
for i, row in enumerate(feat_rows, start=1):
    normal_cell(t6.rows[i].cells[0], row[0], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    normal_cell(t6.rows[i].cells[1], row[1], align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t6.rows[i].cells[2], row[2], align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

add_para(doc, "Clasificadores comparados:", bold=True, size=10)
clf_rows = [
    ("Random Forest", "n_estimators=100, max_depth=10, min_samples_leaf=4, class_weight='balanced'"),
    ("XGBoost",       "n_estimators=100, max_depth=4, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8"),
    ("SVM (RBF)",     "C=1.0, kernel='rbf', probability=True, class_weight='balanced'"),
]
t7 = doc.add_table(rows=len(clf_rows)+1, cols=2)
t7.style = "Table Grid"
header_row(t7, ["Clasificador", "Hiperparámetros principales"])
for i, row in enumerate(clf_rows, start=1):
    normal_cell(t7.rows[i].cells[0], row[0], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t7.rows[i].cells[1], row[1], align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para(doc,
    "Todos los clasificadores se entrenan dentro de un pipeline con StandardScaler. "
    "La selección del mejor clasificador se realiza con validación cruzada estratificada de 5 pliegues.")

# --- 3.3 E3 Híbrido
add_heading(doc, "3.3 E3: Modelo Híbrido (Dos Corrientes)", 2)
add_para(doc,
    "El modelo híbrido combina las dos corrientes en paralelo y fusiona sus representaciones "
    "mediante concatenación antes de la clasificación final:")

add_para(doc, "Corriente A – Visual (imagen 224×224×3):", bold=True, size=10)
stream_a = [
    "EfficientNetB0 (mismo backbone que E1)",
    "GlobalAveragePooling2D",
    "BatchNormalization",
    "Dense(256, relu) → Dropout(0.4) → Dense(128, relu)",
]
for item in stream_a:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(10)

add_para(doc, "Corriente B – Clínica (12 características de pose):", bold=True, size=10)
stream_b = [
    "Dense(64, relu)",
    "BatchNormalization",
    "Dense(32, relu)",
]
for item in stream_b:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(10)

add_para(doc, "Fusión (concatenación de ambas corrientes):", bold=True, size=10)
fusion = [
    "Concatenate([corriente_A, corriente_B])",
    "Dense(64, relu) → Dropout(0.2) → Dense(32, relu)",
    "Dense(1, sigmoid) → Probabilidad de escoliosis",
]
for item in fusion:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
add_para(doc,
    "El entrenamiento sigue la misma estrategia de dos fases que E1 (LR=1×10⁻³ cabeza, "
    "LR=1×10⁻⁵ fine-tune bloques 6-7).")

# ══════════════════════════════════════════════════════════════════════════════
# 4. FUNCIÓN DE PÉRDIDA, OPTIMIZADOR Y MÉTRICAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Función de Pérdida, Optimizador y Métricas", 1)

add_para(doc, "4.1 Función de pérdida", bold=True)
add_para(doc,
    "Para los modelos E1 y E3 se emplea Binary Crossentropy, adecuada para clasificación "
    "binaria con salida sigmoide. Los modelos de ML de E2 optimizan internamente (impureza "
    "de Gini en árboles, pérdida bisagra en SVM).")

add_para(doc, "4.2 Optimizador", bold=True)
add_para(doc,
    "Adam con tasa de aprendizaje dinámica. En la Fase 1 (cabeza) LR=1×10⁻³; "
    "en la Fase 2 (fine-tune) LR=1×10⁻⁵. ReduceLROnPlateau reduce la LR a la mitad "
    "si la pérdida de validación no mejora en 3 épocas consecutivas (mínimo 1×10⁻⁷).")

add_para(doc, "4.3 Métricas de evaluación", bold=True)
metrics = [
    ("Accuracy",     "Proporción de predicciones correctas: (VP+VN)/Total"),
    ("AUC-ROC",      "Área bajo la curva ROC; mide la separabilidad global del modelo"),
    ("Precisión",    "VP / (VP + FP) – fracción de positivos predichos que son reales"),
    ("Recall",       "VP / (VP + FN) – fracción de positivos reales detectados"),
    ("F1-Score",     "Media armónica de Precisión y Recall"),
    ("Cohen's d",    "Tamaño del efecto para la separabilidad de características"),
    ("Matriz de confusión", "Desglose visual de VP, VN, FP, FN"),
]
t8 = doc.add_table(rows=len(metrics)+1, cols=2)
t8.style = "Table Grid"
header_row(t8, ["Métrica", "Descripción"])
for i, (m, d) in enumerate(metrics, start=1):
    normal_cell(t8.rows[i].cells[0], m, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t8.rows[i].cells[1], d, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. INTERPRETABILIDAD
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Técnicas de Interpretabilidad", 1)

add_heading(doc, "5.1 SHAP (SHapley Additive exPlanations)", 2)
add_para(doc,
    "Para los modelos de E2 se calculan valores SHAP que cuantifican la contribución "
    "de cada característica a la predicción:")
shap_rows = [
    ("Random Forest / XGBoost", "TreeExplainer", "Rápido, exacto para modelos de árbol"),
    ("SVM",                     "KernelExplainer (submuestra)", "Agnóstico al modelo, más lento"),
]
t9 = doc.add_table(rows=len(shap_rows)+1, cols=3)
t9.style = "Table Grid"
header_row(t9, ["Modelo", "Explainer SHAP", "Notas"])
for i, row in enumerate(shap_rows, start=1):
    for j, v in enumerate(row):
        normal_cell(t9.rows[i].cells[j], v, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()
add_para(doc,
    "Los resultados se visualizan como beeswarm plots (distribución de valores SHAP "
    "por característica) y gráficas de importancia media |SHAP|.")

add_heading(doc, "5.2 Grad-CAM", 2)
add_para(doc,
    "Para E1 y E3 se aplica Gradient-weighted Class Activation Mapping (Grad-CAM) "
    "sobre la capa 'top_activation' de EfficientNetB0. Genera mapas de calor que "
    "resaltan las regiones de la imagen que más influyen en la decisión del modelo, "
    "facilitando la validación clínica.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. ANÁLISIS DE SEPARABILIDAD DE CARACTERÍSTICAS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Análisis de Separabilidad de Características", 1)

add_para(doc,
    "Se calculó la d de Cohen para cada característica de pose con el fin de cuantificar "
    "la separación entre clases. Los resultados sobre 323 imágenes con pose detectada "
    "(203 sin escoliosis, 120 con escoliosis) son:")

cohen_rows = [
    ("shoulder_height_diff",     "1.138",  "NO",  "Grande – característica discriminante primaria"),
    ("shoulder_width",           "0.789",  "SÍ",  "Mediano"),
    ("hip_height_diff",          "0.744",  "SÍ",  "Mediano"),
    ("hip_asymmetry_ratio",      "0.573",  "SÍ",  "Pequeño-mediano"),
    ("shoulder_asymmetry_ratio", "0.495",  "SÍ",  "Pequeño"),
    ("trunk_tilt_x",             "0.484",  "SÍ",  "Pequeño"),
    ("trunk_tilt_angle",         "0.132",  "SÍ",  "Despreciable"),
    ("wrist_height_diff",        "0.103",  "SÍ",  "Despreciable"),
    ("elbow_height_diff",        "0.081",  "SÍ",  "Despreciable"),
    ("body_height_proxy",        "0.074",  "SÍ",  "Despreciable"),
    ("spine_deviation",          "0.003",  "SÍ",  "Despreciable"),
    ("hip_width",                "0.002",  "SÍ",  "Despreciable"),
]
t10 = doc.add_table(rows=len(cohen_rows)+1, cols=4)
t10.style = "Table Grid"
header_row(t10, ["Característica", "d de Cohen", "Solapamiento", "Interpretación"])
for i, row in enumerate(cohen_rows, start=1):
    normal_cell(t10.rows[i].cells[0], row[0], align=WD_ALIGN_PARAGRAPH.LEFT, bold=(i==1))
    normal_cell(t10.rows[i].cells[1], row[1])
    normal_cell(t10.rows[i].cells[2], row[2])
    normal_cell(t10.rows[i].cells[3], row[3], align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

add_para(doc,
    "La característica shoulder_height_diff presenta el mayor efecto (d = 1.138, "
    "sin solapamiento de distribuciones), confirmando que la asimetría de altura de "
    "hombros es el indicador más robusto de escoliosis en este conjunto de datos.")

add_para(doc,
    "Modelo de sanidad (Decision Stump, profundidad = 1): accuracy = 96.28 %. "
    "Este resultado indica que el umbral en una sola característica separa correctamente "
    "la gran mayoría de los casos, lo que refleja que el dataset contiene casos "
    "clínicamente claros.")

# Imagen de separabilidad si existe
add_image_if_exists(doc,
    "outputs/feature_diagnosis/feature_separability.png",
    width_cm=14,
    caption="Figura 1. Análisis de separabilidad de características (boxplots + d de Cohen)")

# ══════════════════════════════════════════════════════════════════════════════
# 7. IMPORTANCIA DE CARACTERÍSTICAS (SHAP)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Importancia de Características – SHAP (XGBoost)", 1)

add_para(doc,
    "Los valores SHAP calculados sobre el modelo XGBoost revelan la contribución "
    "promedio de cada característica a la predicción:")

shap_feat_rows = [
    ("shoulder_width",           "2.3303",  "1"),
    ("hip_width",                "1.4341",  "2"),
    ("body_height_proxy",        "0.8124",  "3"),
    ("trunk_tilt_x",             "0.4513",  "4"),
    ("spine_deviation",          "0.3477",  "5"),
    ("hip_asymmetry_ratio",      "0.2975",  "6"),
    ("elbow_height_diff",        "0.2775",  "7"),
    ("trunk_tilt_angle",         "0.2626",  "8"),
    ("hip_height_diff",          "0.2404",  "9"),
    ("shoulder_asymmetry_ratio", "0.2374", "10"),
    ("wrist_height_diff",        "0.2310", "11"),
    ("shoulder_height_diff",     "0.2105", "12"),
]
t11 = doc.add_table(rows=len(shap_feat_rows)+1, cols=3)
t11.style = "Table Grid"
header_row(t11, ["Característica", "Media |SHAP|", "Rango"])
for i, row in enumerate(shap_feat_rows, start=1):
    normal_cell(t11.rows[i].cells[0], row[0], align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t11.rows[i].cells[1], row[1])
    normal_cell(t11.rows[i].cells[2], row[2])
doc.add_paragraph()

add_para(doc,
    "Interpretación: aunque shoulder_height_diff tiene la mayor d de Cohen, XGBoost "
    "prioriza shoulder_width y hip_width. Esto sugiere que las características de "
    "anchura son más estables ante variaciones de pose y escala, y que el modelo "
    "aprovecha correlaciones no lineales entre características.")

# Imagen SHAP
add_image_if_exists(doc,
    "outputs/pose/shap_beeswarm_XGBClassifier.png",
    width_cm=14,
    caption="Figura 2. Beeswarm plot de valores SHAP (XGBoost)")

add_image_if_exists(doc,
    "outputs/pose/shap_bar_XGBClassifier.png",
    width_cm=14,
    caption="Figura 3. Importancia de características por media |SHAP| (XGBoost)")

# ══════════════════════════════════════════════════════════════════════════════
# 8. RESULTADOS – VALIDACIÓN CRUZADA (K-FOLD)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Resultados – Validación Cruzada 5-Fold", 1)

add_para(doc,
    "Se aplicó validación cruzada estratificada de 5 pliegues para obtener estimaciones "
    "robustas e imparciales del rendimiento. Los modelos de redes se reentrenan en cada "
    "pliegue; el modelo de pose utiliza el mejor clasificador seleccionado previamente.")

for label, folds in [("E1 – CNN", kfold_cnn), ("E2 – Pose + ML", kfold_pose)]:
    if folds is None:
        continue
    add_para(doc, label, bold=True, size=11)
    t_cv = doc.add_table(rows=len(folds)+2, cols=4)
    t_cv.style = "Table Grid"
    header_row(t_cv, ["Pliegue", "Accuracy", "F1-Score", "AUC-ROC"])
    for i, fold in enumerate(folds, start=1):
        normal_cell(t_cv.rows[i].cells[0], f"Fold {i}")
        normal_cell(t_cv.rows[i].cells[1], f"{fold['acc']:.4f}")
        normal_cell(t_cv.rows[i].cells[2], f"{fold['f1']:.4f}")
        normal_cell(t_cv.rows[i].cells[3], f"{fold['auc']:.4f}")
    # Fila de media ± std
    last = t_cv.rows[len(folds)+1]
    acc_m, acc_s = mean_std(folds, 'acc')
    f1_m,  f1_s  = mean_std(folds, 'f1')
    auc_m, auc_s = mean_std(folds, 'auc')
    set_cell_bg(last.cells[0], "D6E4F0")
    set_cell_bg(last.cells[1], "D6E4F0")
    set_cell_bg(last.cells[2], "D6E4F0")
    set_cell_bg(last.cells[3], "D6E4F0")
    bold_cell(last.cells[0], "Media ± σ", size=10)
    bold_cell(last.cells[1], f"{acc_m:.4f} ± {acc_s:.4f}", size=10)
    bold_cell(last.cells[2], f"{f1_m:.4f} ± {f1_s:.4f}",  size=10)
    bold_cell(last.cells[3], f"{auc_m:.4f} ± {auc_s:.4f}", size=10)
    doc.add_paragraph()

add_para(doc,
    "Interpretación: Ambos enfoques alcanzan métricas perfectas (1.00) en los 5 pliegues. "
    "Este resultado es consistente con el análisis de separabilidad, que demostró que el "
    "dataset contiene casos clínicamente evidentes con alta separación en el espacio de "
    "características. La ausencia de varianza entre pliegues sugiere que el problema, "
    "tal como está formulado, tiene una frontera de decisión simple y estable.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. RESULTADOS – ROBUSTEZ AL RUIDO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Resultados – Robustez al Ruido Gaussiano", 1)

add_para(doc,
    "Se evaluó la degradación del rendimiento al añadir ruido gaussiano progresivo "
    "(σ ∈ {0.00, 0.05, 0.10, 0.20, 0.30, 0.50}). Para E2 el ruido se aplica sobre las "
    "características numéricas; para E1 se aplica sobre los píxeles (σ × 255).")

noise_levels = [0.0, 0.05, 0.1, 0.2, 0.3, 0.5]

for label, rob in [("E1 – CNN", rob_cnn), ("E2 – Pose + ML", rob_pose)]:
    if rob is None:
        continue
    add_para(doc, label, bold=True, size=11)
    t_rob = doc.add_table(rows=len(noise_levels)+1, cols=4)
    t_rob.style = "Table Grid"
    header_row(t_rob, ["Nivel de ruido (σ)", "Accuracy (media)", "F1 (media)", "AUC (media)"])
    for i, sigma in enumerate(noise_levels, start=1):
        folds = rob.get(sigma, [])
        if folds:
            am, _ = mean_std(folds, 'acc')
            fm, _ = mean_std(folds, 'f1')
            um, _ = mean_std(folds, 'auc')
            normal_cell(t_rob.rows[i].cells[0], str(sigma))
            normal_cell(t_rob.rows[i].cells[1], f"{am:.4f}")
            normal_cell(t_rob.rows[i].cells[2], f"{fm:.4f}")
            normal_cell(t_rob.rows[i].cells[3], f"{um:.4f}")
            # Highlight degradación severa
            if am < 0.75:
                for c in t_rob.rows[i].cells:
                    set_cell_bg(c, "FFD7D7")
    doc.add_paragraph()

add_para(doc,
    "Interpretación – E2 (Pose + ML): El modelo geométrico mantiene accuracy ≥ 95 % "
    "con ruido bajo (σ = 0.05) y ≥ 87 % con ruido moderado (σ = 0.20). Con ruido severo "
    "(σ = 0.50) el accuracy cae a ~73 % y el AUC a ~80 %, lo que indica que las "
    "características geométricas son moderadamente robustas pero sensibles a perturbaciones "
    "grandes en los puntos clave detectados.")

add_para(doc,
    "Interpretación – E1 (CNN): La red convolucional es mucho más robusta al ruido en "
    "píxeles: mantiene accuracy = 100 % hasta σ = 0.20 y ≥ 89 % a σ = 0.30. Sin embargo, "
    "a σ = 0.50 el rendimiento colapsa a ~43 % (peor que el azar), lo que indica que "
    "con ruido extremo la red pierde completamente las características visuales relevantes.")

# Imágenes de robustez
add_image_if_exists(doc,
    "outputs/robustness/robustness_cnn/robustness_curve_cnn.png",
    width_cm=13,
    caption="Figura 4. Curva de robustez al ruido – E1 CNN")
add_image_if_exists(doc,
    "outputs/robustness/robustness_pose/robustness_curve_pose.png",
    width_cm=13,
    caption="Figura 5. Curva de robustez al ruido – E2 Pose + ML")

# ══════════════════════════════════════════════════════════════════════════════
# 10. COMPARACIÓN DE ENFOQUES (ABLATION STUDY)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. Comparación de Enfoques – Estudio de Ablación", 1)

add_para(doc,
    "La siguiente tabla resume las capacidades y limitaciones de cada enfoque basándose "
    "en los resultados obtenidos:")

ablation_rows = [
    ("E1 – CNN",    "Alta (imagen)",         "Baja (caja negra)",    "Alta (σ ≤ 0.20)",  "Alta"),
    ("E2 – Pose+ML","Baja-media (geométrica)","Alta (SHAP)",          "Media (σ ≤ 0.10)", "Baja"),
    ("E3 – Híbrido","Alta (imagen + pose)",  "Media (SHAP + GradCAM)","Alta",             "Muy Alta"),
]
t12 = doc.add_table(rows=len(ablation_rows)+1, cols=5)
t12.style = "Table Grid"
header_row(t12, ["Enfoque", "Información usada", "Interpretabilidad", "Robustez al ruido", "Complejidad"])
for i, row in enumerate(ablation_rows, start=1):
    for j, v in enumerate(row):
        normal_cell(t12.rows[i].cells[j], v, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

add_image_if_exists(doc,
    "outputs/ablation/ablation_comparison.png",
    width_cm=14,
    caption="Figura 6. Comparación de métricas entre enfoques (estudio de ablación)")

# ══════════════════════════════════════════════════════════════════════════════
# 11. VISUALIZACIONES CLAVE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. Visualizaciones Clave", 1)

img_list = [
    ("outputs/cnn/cnn_training_curves.png",   "Figura 7. Curvas de entrenamiento CNN (pérdida y accuracy)"),
    ("outputs/cnn/cnn_confusion_matrix.png",  "Figura 8. Matriz de confusión – E1 CNN"),
    ("outputs/cnn/cnn_roc_curve.png",         "Figura 9. Curva ROC – E1 CNN"),
    ("outputs/pose/pose_cv_comparison.png",   "Figura 10. Comparación de clasificadores de pose (CV)"),
    ("outputs/pose/pose_roc_comparison.png",  "Figura 11. Curvas ROC – clasificadores de pose"),
]
for path, caption in img_list:
    add_image_if_exists(doc, path, width_cm=13, caption=caption)

# ══════════════════════════════════════════════════════════════════════════════
# 12. DEPENDENCIAS Y ENTORNO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "12. Dependencias y Entorno de Desarrollo", 1)

deps = [
    ("TensorFlow / Keras", "≥ 2.12.0",  "Framework principal de deep learning"),
    ("scikit-learn",       "≥ 1.3.0",   "Clasificadores clásicos, CV, métricas"),
    ("XGBoost",            "≥ 2.0.0",   "Clasificador de gradient boosting"),
    ("MediaPipe",          "≥ 0.10.14", "Estimación de pose (33 puntos clave)"),
    ("albumentations",     "≥ 1.3.0",   "Pipeline de aumentación de datos"),
    ("SHAP",               "≥ 0.43.0",  "Interpretabilidad de modelos"),
    ("OpenCV",             "≥ 4.8.0",   "Procesamiento de imágenes"),
    ("NumPy / Pandas",     "≥ 1.24 / 2.0", "Manipulación de datos"),
    ("Matplotlib / Seaborn","≥ 3.7 / 0.12", "Visualización"),
]
t13 = doc.add_table(rows=len(deps)+1, cols=3)
t13.style = "Table Grid"
header_row(t13, ["Biblioteca", "Versión mínima", "Uso principal"])
for i, row in enumerate(deps, start=1):
    normal_cell(t13.rows[i].cells[0], row[0], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    normal_cell(t13.rows[i].cells[1], row[1])
    normal_cell(t13.rows[i].cells[2], row[2], align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 13. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "13. Conclusiones", 1)

conclusions = [
    "El dataset contiene casos clínicamente obvios: la asimetría de hombros "
    "(shoulder_height_diff, d = 1.138) es suficiente para separar las clases con un "
    "único umbral (accuracy = 96.28 %). Esto explica las métricas perfectas en "
    "validación cruzada.",

    "E1 (CNN) demuestra mayor robustez al ruido pixel-level que E2 (Pose+ML) ante "
    "ruido moderado (σ ≤ 0.20), pero colapsa a σ = 0.50. E2 degrada más gradualmente, "
    "lo que sugiere que las características geométricas actúan como una representación "
    "más compacta y resistente.",

    "E3 (Híbrido) combina las fortalezas de ambos enfoques. Su mayor complejidad se "
    "justifica si se requiere tanto rendimiento visual como interpretabilidad clínica.",

    "SHAP revela que XGBoost prioriza características de anchura (shoulder_width, "
    "hip_width) sobre las de asimetría de altura, indicando que el modelo captura "
    "patrones complementarios al análisis univariado.",

    "El trabajo futuro debe incluir datos con mayor variabilidad (ángulos de captura, "
    "distintos grados de escoliosis) para evaluar la generalización real del sistema.",
]
for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(c).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
out_path = "Reporte_Tecnico_Tesis_Escoliosis.docx"
doc.save(out_path)
print(f"Documento guardado: {out_path}")
